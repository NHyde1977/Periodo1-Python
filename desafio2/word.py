from docx import Document
from docx.shared import Pt

doc = Document()

def add_heading(text, level):
    doc.add_heading(text, level=level)

def add_paragraph(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    p.paragraph_format.space_after = Pt(8)
    return p

# Título
add_heading("Gestão de Projeto e Controle da Evolução de Software em Equipes de Desenvolvimento", 0)

# Contexto
add_heading("Contexto da Demanda", 1)
add_paragraph("Você está liderando um time de desenvolvimento com 30 colaboradores, divididos em múltiplos projetos. "
              "Cada projeto exige controle eficaz da evolução do software, incluindo gestão de código, operação e entrega. "
              "O objetivo é estruturar os times e processos de forma que esses elementos aumentem a produtividade geral, "
              "reduzam falhas e melhorem a comunicação entre os envolvidos.")

# Fraquezas
add_heading("Principais Fraquezas Identificadas", 1)
fraquezas = [
    "Falta de documentação ou documentação inadequada.",
    "Presença de soluções temporárias ('gambiarras') que se tornam permanentes.",
    "Uso de versões desatualizadas de serviços e bibliotecas.",
    "Turnover de desenvolvedores sem processos de handover claros.",
    "Falta de visão holística e integração entre os projetos.",
    "Uso de tecnologias que não se comunicam entre si.",
    "Disputa entre times por recursos escassos (infraestrutura, habilidades, licenças etc.).",
    "Gargalos na esteira de produção: testes, correções e janelas de atualização.",
    "Histórias e requisitos mal definidos ou ambíguos.",
    "Foco excessivo na entrega de novas funcionalidades, negligenciando correções.",
    "Ausência de padrão entre as equipes nos processos de backlog, priorização, desenvolvimento, testes e produção."
]
for f in fraquezas:
    add_paragraph(f)

# Soluções
add_heading("Propostas de Solução", 1)
solucoes = {
    "1. Melhoria na Documentação": [
        "Estabelecer a documentação como parte da Definition of Done (DoD).",
        "Designar um responsável por garantir a documentação ao final de cada tarefa.",
        "Estimular o uso de ferramentas como Confluence, Notion ou GitBook para centralizar o conhecimento técnico."
    ],
    "2 e 3. Gestão de Dívida Técnica e Atualizações": [
        "Reservar tempo dedicado em cada sprint para manutenção, correções e refatorações.",
        "Criar um cronograma de atualização tecnológica dos serviços e bibliotecas.",
        "Considerar a contratação de equipes especializadas para tarefas específicas de modernização.",
        "Incluir a dívida técnica no backlog como item prioritário quando houver risco operacional."
    ],
    "4. Turnover e Retenção de Conhecimento": [
        "Incentivar uma cultura contínua de documentação e compartilhamento de conhecimento.",
        "Criar checklists de offboarding técnico (credenciais, pendências, walkthroughs).",
        "Formalizar cláusulas contratuais que assegurem o handover adequado antes da saída de profissionais-chave.",
        "Atuar com tech leads como multiplicadores de boas práticas."
    ],
    "5. Integração e Visão Sistêmica": [
        "Realizar reuniões quinzenais ou mensais do tipo “all hands” para atualizar o time sobre o panorama geral.",
        "Usar ferramentas como Jira Portfolio, Miro ou dashboards visuais para mapear como os projetos se inter-relacionam.",
        "Compartilhar roadmaps de produto de forma transparente."
    ],
    "6. Integração Tecnológica": [
        "Promover o uso de arquiteturas reutilizáveis e escaláveis, como microsserviços e APIs RESTful.",
        "Estabelecer padrões de comunicação entre sistemas (ex: OpenAPI, GraphQL).",
        "Criar um comitê de governança técnica para avaliar decisões de arquitetura e tecnologia."
    ],
    "7. Gerenciamento de Recursos Escassos": [
        "Realizar revisões periódicas de orçamento e capacidade.",
        "Utilizar modelos de priorização (WSJF, RICE, custo/benefício) para alocar recursos.",
        "Permitir a substituição de entregáveis que não dependam diretamente dos recursos indisponíveis.",
        "Automatizar o máximo possível para reduzir dependências manuais."
    ],
    "8. Melhoria da Esteira de Produção": [
        "Aplicar Value Stream Mapping para identificar e eliminar gargalos.",
        "Implementar ferramentas de CI/CD (ex: GitHub Actions, GitLab CI, Jenkins).",
        "Automatizar testes e validações (ex: testes unitários, testes de integração e linting).",
        "Medir métricas como cycle time, lead time, failure rate e mean time to recovery (MTTR)."
    ],
    "9 e 11. Padrões de Requisitos e Processos": [
        "Criar um Guia de Desenvolvimento e Boas Práticas para padronizar escrita de histórias, testes e entregas.",
        "Definir responsáveis por revisar e organizar os cards e documentação do backlog.",
        "Promover treinamentos e workshops regulares com foco em agilidade, análise de requisitos e testes.",
        "Adotar um modelo unificado de workflow entre as equipes (ex: Kanban, Scrum, Scrumban)."
    ],
    "10. Equilíbrio entre Funcionalidades e Correções": [
        "Utilizar critérios como valor gerado, impacto estratégico e custo de oportunidade para priorizar entregas.",
        "Criar dashboards de bugs, débitos técnicos e métricas de qualidade para visibilidade.",
        "Estabelecer SLAs internos para resolução de falhas e bugs críticos."
    ]
}
for titulo, itens in solucoes.items():
    add_heading(titulo, 2)
    for item in itens:
        add_paragraph(item)

add_heading("Conclusão", 1)
add_paragraph("A gestão eficaz da evolução de software requer uma combinação equilibrada entre processos bem definidos, "
              "cultura de colaboração, uso de boas práticas e tecnologias adequadas. A implementação das ações propostas "
              "permitirá maior alinhamento entre as equipes, controle mais eficiente do progresso dos projetos.")
